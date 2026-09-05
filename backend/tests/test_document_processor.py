"""Unit tests for document processing service.

Tests PDF/DOCX parsing, file validation, security checks, and deduplication.
"""

from __future__ import annotations

import uuid

import pytest
from fastapi import HTTPException, status
from sqlalchemy import select

from app.auth.models import User
from app.modules.documents.models import CandidateDocument, DocumentJob
from app.modules.documents.service import DocumentService
from app.services.document_processor import DocumentProcessingError, DocumentProcessor
from app.storage.document_storage import (
    MAX_FILE_SIZE_BYTES,
    DocumentStorageClient,
    DocumentStorageError,
    compute_file_hash,
    validate_file_size,
    validate_mime_type,
)


@pytest.fixture
def processor() -> DocumentProcessor:
    """Create document processor instance."""
    return DocumentProcessor()


@pytest.fixture
def storage_client() -> DocumentStorageClient:
    """Create document storage client instance."""
    return DocumentStorageClient()


@pytest.fixture
def sample_pdf_data() -> bytes:
    """Create minimal valid PDF data for testing."""
    # Minimal PDF structure
    pdf_content = b"""%PDF-1.4
1 0 obj
<<
/Type /Catalog
/Pages 2 0 R
>>
endobj
2 0 obj
<<
/Type /Pages
/Kids [3 0 R]
/Count 1
>>
endobj
3 0 obj
<<
/Type /Page
/Parent 2 0 R
/Resources <<
/Font <<
/F1 <<
/Type /Font
/Subtype /Type1
/BaseFont /Helvetica
>>
>>
>>
/MediaBox [0 0 612 792]
/Contents 4 0 R
>>
endobj
4 0 obj
<<
/Length 44
>>
stream
BT
/F1 12 Tf
100 700 Td
(Test PDF Content) Tj
ET
endstream
endobj
xref
0 5
0000000000 65535 f
0000000009 00000 n
0000000058 00000 n
0000000115 00000 n
0000000317 00000 n
trailer
<<
/Size 5
/Root 1 0 R
>>
startxref
410
%%EOF
"""
    return pdf_content


@pytest.fixture
def sample_docx_data() -> bytes:
    """Create minimal valid DOCX data for testing."""
    try:
        from io import BytesIO

        import docx

        doc = docx.Document()
        doc.add_paragraph("Test DOCX Content")
        doc.add_paragraph("This is a sample resume.")

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()
    except ImportError:
        pytest.skip("python-docx not installed")


# ── Document Processing Tests ──────────────────────────────────────


def test_extract_pdf_text(processor: DocumentProcessor, sample_pdf_data: bytes) -> None:
    """Test PDF text extraction with valid file."""
    result = processor.extract_pdf_text(sample_pdf_data)

    assert "text" in result
    assert "Test PDF Content" in result["text"]
    assert result["page_count"] == 1
    assert result["token_count"] > 0
    assert "metadata" in result


def test_extract_pdf_text_corrupted_file(processor: DocumentProcessor) -> None:
    """Test PDF extraction fails with corrupted file."""
    corrupted_pdf = b"Not a PDF file"

    with pytest.raises(DocumentProcessingError, match="Failed to process PDF"):
        processor.extract_pdf_text(corrupted_pdf)


def test_extract_pdf_text_empty_file(processor: DocumentProcessor) -> None:
    """Test PDF extraction fails with empty file."""
    empty_pdf = b""

    with pytest.raises(DocumentProcessingError):
        processor.extract_pdf_text(empty_pdf)


def test_extract_docx_text(processor: DocumentProcessor, sample_docx_data: bytes) -> None:
    """Test DOCX text extraction with valid file."""
    result = processor.extract_docx_text(sample_docx_data)

    assert "text" in result
    assert "Test DOCX Content" in result["text"]
    assert "sample resume" in result["text"]
    assert result["paragraph_count"] >= 2
    assert result["token_count"] > 0
    assert "metadata" in result


def test_extract_docx_text_corrupted_file(processor: DocumentProcessor) -> None:
    """Test DOCX extraction fails with corrupted file."""
    corrupted_docx = b"Not a DOCX file"

    with pytest.raises(DocumentProcessingError, match="Failed to process DOCX"):
        processor.extract_docx_text(corrupted_docx)


def test_process_document_pdf(processor: DocumentProcessor, sample_pdf_data: bytes) -> None:
    """Test document processing routes to PDF handler."""
    result = processor.process_document(sample_pdf_data, "application/pdf")

    assert "text" in result
    assert "page_count" in result


def test_process_document_docx(processor: DocumentProcessor, sample_docx_data: bytes) -> None:
    """Test document processing routes to DOCX handler."""
    result = processor.process_document(
        sample_docx_data,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    assert "text" in result
    assert "paragraph_count" in result


def test_process_document_unsupported_type(processor: DocumentProcessor) -> None:
    """Test document processing fails with unsupported MIME type."""
    with pytest.raises(DocumentProcessingError, match="Unsupported MIME type"):
        processor.process_document(b"test", "application/unknown")


# ── Storage Validation Tests ───────────────────────────────────────


def test_validate_file_size_valid() -> None:
    """Test file size validation passes for valid size."""
    validate_file_size(1024)  # 1KB
    validate_file_size(MAX_FILE_SIZE_BYTES)  # Max size
    # Should not raise


def test_validate_file_size_exceeds_limit() -> None:
    """Test file size validation fails for oversized file."""
    with pytest.raises(DocumentStorageError, match="exceeds maximum"):
        validate_file_size(MAX_FILE_SIZE_BYTES + 1)


def test_validate_mime_type_pdf() -> None:
    """Test MIME type validation for PDF."""
    ext = validate_mime_type("application/pdf")
    assert ext == "pdf"


def test_validate_mime_type_docx() -> None:
    """Test MIME type validation for DOCX."""
    ext = validate_mime_type(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert ext == "docx"


def test_validate_mime_type_with_charset() -> None:
    """Test MIME type validation strips charset."""
    ext = validate_mime_type("application/pdf; charset=utf-8")
    assert ext == "pdf"


def test_validate_mime_type_unsupported() -> None:
    """Test MIME type validation fails for unsupported type."""
    with pytest.raises(DocumentStorageError, match="Unsupported file type"):
        validate_mime_type("application/x-msword")  # Old .doc format


def test_compute_file_hash() -> None:
    """Test file hash computation for deduplication."""
    data1 = b"test document content"
    data2 = b"test document content"
    data3 = b"different content"

    hash1 = compute_file_hash(data1)
    hash2 = compute_file_hash(data2)
    hash3 = compute_file_hash(data3)

    # Same content produces same hash
    assert hash1 == hash2
    assert len(hash1) == 64  # SHA256 hex length

    # Different content produces different hash
    assert hash1 != hash3


# ── Storage Upload Tests ───────────────────────────────────────────


@pytest.mark.asyncio
async def test_upload_document_valid_pdf(
    storage_client: DocumentStorageClient,
    sample_pdf_data: bytes,
) -> None:
    """Test document upload with valid PDF."""
    storage_path, file_hash, file_size = await storage_client.upload_document(
        sample_pdf_data,
        "resume.pdf",
        "application/pdf",
        "user-123",
        "cv",
    )

    assert storage_path.startswith("documents/user-123/cv/")
    assert storage_path.endswith(".pdf")
    assert len(file_hash) == 64
    assert file_size == len(sample_pdf_data)


@pytest.mark.asyncio
async def test_upload_document_valid_docx(
    storage_client: DocumentStorageClient,
    sample_docx_data: bytes,
) -> None:
    """Test document upload with valid DOCX."""
    storage_path, file_hash, file_size = await storage_client.upload_document(
        sample_docx_data,
        "resume.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "user-456",
        "cv",
    )

    assert storage_path.startswith("documents/user-456/cv/")
    assert storage_path.endswith(".docx")
    assert len(file_hash) == 64
    assert file_size == len(sample_docx_data)


@pytest.mark.asyncio
async def test_upload_document_exceeds_size_limit(
    storage_client: DocumentStorageClient,
) -> None:
    """Test document upload fails when file exceeds size limit."""
    oversized_data = b"x" * (MAX_FILE_SIZE_BYTES + 1)

    with pytest.raises(DocumentStorageError, match="exceeds maximum"):
        await storage_client.upload_document(
            oversized_data,
            "huge.pdf",
            "application/pdf",
            "user-789",
            "cv",
        )


@pytest.mark.asyncio
async def test_upload_document_unsupported_type(
    storage_client: DocumentStorageClient,
) -> None:
    """Test document upload fails with unsupported MIME type."""
    with pytest.raises(DocumentStorageError, match="Unsupported file type"):
        await storage_client.upload_document(
            b"test",
            "document.txt",
            "text/plain",
            "user-999",
            "cv",
        )


@pytest.mark.asyncio
async def test_upload_document_corrupted_file(
    storage_client: DocumentStorageClient,
) -> None:
    """Test document upload fails with corrupted/empty file."""
    tiny_file = b"x" * 50  # Less than 100 bytes

    with pytest.raises(DocumentStorageError, match="corrupted or empty"):
        await storage_client.upload_document(
            tiny_file,
            "corrupted.pdf",
            "application/pdf",
            "user-111",
            "cv",
        )


@pytest.mark.asyncio
async def test_duplicate_upload_different_hash(
    storage_client: DocumentStorageClient,
    sample_pdf_data: bytes,
) -> None:
    """Test uploading same file produces same hash (for deduplication)."""
    # Upload same file twice
    path1, hash1, size1 = await storage_client.upload_document(
        sample_pdf_data,
        "resume1.pdf",
        "application/pdf",
        "user-222",
        "cv",
    )

    path2, hash2, size2 = await storage_client.upload_document(
        sample_pdf_data,
        "resume2.pdf",
        "application/pdf",
        "user-222",
        "cv",
    )

    # Different storage paths (unique IDs)
    assert path1 != path2

    # But same file hash (for deduplication detection)
    assert hash1 == hash2
    assert size1 == size2


# ── Storage Download Tests ─────────────────────────────────────────


@pytest.mark.asyncio
async def test_upload_then_download_round_trip(
    storage_client: DocumentStorageClient,
    sample_pdf_data: bytes,
) -> None:
    """Bytes written by upload_document must be retrievable via
    download_document, exercising the local-cache fallback used in tests
    (no R2 credentials configured)."""
    storage_path, _, _ = await storage_client.upload_document(
        sample_pdf_data,
        "resume.pdf",
        "application/pdf",
        "user-round-trip",
        "cv",
    )

    downloaded = await storage_client.download_document(storage_path)

    assert downloaded == sample_pdf_data


@pytest.mark.asyncio
async def test_download_document_missing_object_raises(
    storage_client: DocumentStorageClient,
) -> None:
    """Downloading a storage path that was never written raises a clear
    DocumentStorageError instead of propagating a raw filesystem error."""
    with pytest.raises(DocumentStorageError, match="Download failed"):
        await storage_client.download_document("documents/user-missing/cv/does-not-exist.pdf")


# ── reprocess_document Tests ───────────────────────────────────────


async def _create_reprocess_test_user(db) -> User:
    user = User(
        email=f"reprocess-{uuid.uuid4().hex[:10]}@example.com",
        first_name="Reprocess",
        last_name="Tester",
        is_verified=True,
    )
    db.add(user)
    await db.commit()
    await db.refresh(user)
    return user


async def _create_reprocess_test_document(db, user_id, **overrides) -> CandidateDocument:
    fields = {
        "user_id": user_id,
        "document_type": "cv",
        "original_filename": "resume.pdf",
        "storage_path": f"documents/{user_id}/cv/{uuid.uuid4().hex}.pdf",
        "file_hash": uuid.uuid4().hex,
        "file_size_bytes": 2048,
        "processing_status": "completed",
    }
    fields.update(overrides)
    doc = CandidateDocument(**fields)
    db.add(doc)
    await db.commit()
    await db.refresh(doc)
    return doc


@pytest.mark.asyncio
async def test_reprocess_document_legacy_row_without_mime_type_raises(db) -> None:
    """Legacy rows (mime_type=None, predating the storage-backed upload path)
    can't be reprocessed automatically since there are no real stored bytes
    to re-extract from — the service must raise a clear error instead of
    silently no-op'ing."""
    user = await _create_reprocess_test_user(db)
    doc = await _create_reprocess_test_document(db, user.id, mime_type=None)

    service = DocumentService(db)

    with pytest.raises(HTTPException) as exc_info:
        await service.reprocess_document(str(doc.id), user.id)

    assert exc_info.value.status_code == status.HTTP_409_CONFLICT
    assert "re-upload" in exc_info.value.detail


@pytest.mark.asyncio
async def test_reprocess_document_success_enqueues_job(db) -> None:
    """A document with mime_type + real stored bytes gets its status/text
    reset and a new DocumentJob enqueued via the same RQ wiring as
    upload_document."""
    user = await _create_reprocess_test_user(db)

    storage_client = DocumentStorageClient()
    storage_path, file_hash, file_size = await storage_client.upload_document(
        b"%PDF-1.4 " + b"x" * 200,
        "resume.pdf",
        "application/pdf",
        str(user.id),
        "cv",
    )

    doc = await _create_reprocess_test_document(
        db,
        user.id,
        storage_path=storage_path,
        mime_type="application/pdf",
        file_hash=file_hash,
        file_size_bytes=file_size,
        processing_status="failed",
        raw_text="stale text",
        extracted_data={"stale": True},
    )

    service = DocumentService(db)
    response = await service.reprocess_document(str(doc.id), user.id)

    assert response.document_id == str(doc.id)
    assert response.job_id

    await db.refresh(doc)
    assert doc.processing_status == "pending"
    assert doc.raw_text is None
    assert doc.extracted_data is None

    result = await db.execute(
        select(DocumentJob).where(
            DocumentJob.document_id == doc.id,
            DocumentJob.job_type == "reprocess",
        )
    )
    job = result.scalar_one()
    assert str(job.id) == response.job_id


@pytest.mark.asyncio
async def test_reprocess_document_missing_stored_bytes_raises(db) -> None:
    """If the document row has a mime_type but its bytes are gone from
    storage (e.g. deleted out-of-band), reprocess must surface a clear
    404 instead of an unhandled storage exception."""
    user = await _create_reprocess_test_user(db)
    doc = await _create_reprocess_test_document(
        db,
        user.id,
        mime_type="application/pdf",
        storage_path=f"documents/{user.id}/cv/never-uploaded.pdf",
    )

    service = DocumentService(db)

    with pytest.raises(HTTPException) as exc_info:
        await service.reprocess_document(str(doc.id), user.id)

    assert exc_info.value.status_code == status.HTTP_404_NOT_FOUND


# ── Coverage Summary ───────────────────────────────────────────────
# This test suite covers:
# - PDF parsing (valid, corrupted, empty)
# - DOCX parsing (valid, corrupted)
# - File validation (size, MIME type, corrupted)
# - Security checks (size limits, type validation)
# - Deduplication (file hashing)
# - Storage operations (upload, path generation)
# - Error handling for all failure modes
#
# Expected coverage: >80% for:
# - app/services/document_processor.py
# - app/storage/document_storage.py
