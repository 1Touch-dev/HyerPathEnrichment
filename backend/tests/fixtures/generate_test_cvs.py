"""Generate test CV fixtures for document processing tests."""

from pathlib import Path
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document

FIXTURES_DIR = Path(__file__).parent


def generate_sample_cv_pdf():
    """Generate a standard 2-page CV PDF."""
    pdf_path = FIXTURES_DIR / "sample_cv.pdf"

    c = canvas.Canvas(str(pdf_path), pagesize=letter)
    width, height = letter

    # Page 1
    y = height - 50
    c.setFont("Helvetica-Bold", 16)
    c.drawString(50, y, "John Doe")

    y -= 20
    c.setFont("Helvetica", 12)
    c.drawString(50, y, "Senior Backend Engineer")

    y -= 30
    c.drawString(50, y, "Email: john.doe@example.com")
    y -= 20
    c.drawString(50, y, "Phone: +91-98765-43210")
    y -= 20
    c.drawString(50, y, "LinkedIn: linkedin.com/in/johndoe")

    y -= 40
    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, y, "PROFESSIONAL SUMMARY")

    y -= 25
    c.setFont("Helvetica", 11)
    text = "Experienced software engineer with 7+ years building scalable APIs"
    c.drawString(50, y, text)
    y -= 15
    text = "and microservices. Proficient in Python, FastAPI, PostgreSQL, Docker."
    c.drawString(50, y, text)

    y -= 35
    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, y, "EXPERIENCE")

    y -= 25
    c.setFont("Helvetica-Bold", 12)
    c.drawString(50, y, "Senior Backend Engineer at TechCorp (2020-2024)")

    y -= 20
    c.setFont("Helvetica", 11)
    c.drawString(60, y, "- Built payment processing API handling 50M requests/day")
    y -= 15
    c.drawString(60, y, "- Led team of 4 engineers")
    y -= 15
    c.drawString(60, y, "- Reduced latency by 60% using Redis caching")

    # Page 2
    c.showPage()
    y = height - 50

    c.setFont("Helvetica-Bold", 12)
    c.drawString(50, y, "Backend Developer at StartupXYZ (2017-2020)")

    y -= 20
    c.setFont("Helvetica", 11)
    c.drawString(60, y, "- Developed REST APIs using Flask and PostgreSQL")
    y -= 15
    c.drawString(60, y, "- Implemented CI/CD pipeline with Jenkins")

    y -= 35
    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, y, "EDUCATION")

    y -= 25
    c.setFont("Helvetica", 11)
    c.drawString(50, y, "B.Tech Computer Science, IIT Delhi (2013-2017)")

    y -= 35
    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, y, "SKILLS")

    y -= 25
    c.setFont("Helvetica", 11)
    c.drawString(
        50, y, "Technical: Python, FastAPI, Flask, PostgreSQL, Redis, Docker, Kubernetes, AWS"
    )
    y -= 15
    c.drawString(50, y, "Soft Skills: Team Leadership, Problem Solving, Communication")
    y -= 15
    c.drawString(50, y, "Languages: English (Fluent), Hindi (Native)")

    c.save()
    print(f"Created: {pdf_path}")


def generate_minimal_cv_pdf():
    """Generate a CV with ~70% completeness (missing fields)."""
    pdf_path = FIXTURES_DIR / "sample_cv_minimal.pdf"

    c = canvas.Canvas(str(pdf_path), pagesize=letter)
    width, height = letter

    y = height - 50
    c.setFont("Helvetica-Bold", 16)
    c.drawString(50, y, "Jane Smith")

    y -= 30
    c.setFont("Helvetica", 12)
    c.drawString(50, y, "Email: jane@example.com")
    # No phone, no LinkedIn

    y -= 40
    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, y, "EXPERIENCE")

    y -= 25
    c.setFont("Helvetica-Bold", 12)
    c.drawString(50, y, "Software Engineer at CompanyABC")
    # No dates, vague

    y -= 20
    c.setFont("Helvetica", 11)
    c.drawString(60, y, "- Worked on backend systems")
    # No metrics, no details

    y -= 35
    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, y, "SKILLS")

    y -= 25
    c.setFont("Helvetica", 11)
    c.drawString(50, y, "Python, JavaScript")
    # Minimal skills list

    c.save()
    print(f"Created: {pdf_path}")


def generate_sample_cv_docx():
    """Generate a DOCX format CV."""
    docx_path = FIXTURES_DIR / "sample_cv.docx"

    doc = Document()

    doc.add_heading("Sarah Chen", level=1)
    doc.add_paragraph("ML Engineer | sarah.chen@example.com")

    doc.add_heading("Professional Summary", level=2)
    doc.add_paragraph(
        "Machine Learning Engineer with 5+ years building recommendation systems "
        "at Netflix and Amazon. Expert in NLP and deep learning."
    )

    doc.add_heading("Experience", level=2)
    doc.add_paragraph("ML Engineer at Amazon (2021-2024)", style="Heading 3")
    doc.add_paragraph("- Built recommendation engine serving 100M users")
    doc.add_paragraph("- Improved click-through rate by 25% using transformers")

    doc.add_heading("Skills", level=2)
    doc.add_paragraph("Python, TensorFlow, PyTorch, AWS, Kubernetes")

    doc.save(docx_path)
    print(f"Created: {docx_path}")


def generate_malformed_pdf():
    """Generate a corrupted PDF for error testing."""
    pdf_path = FIXTURES_DIR / "malformed.pdf"

    # Write invalid PDF content
    with open(pdf_path, "wb") as f:
        f.write(b"%PDF-1.4\n")
        f.write(b"This is not valid PDF content!\n")
        f.write(b"Random bytes: \x00\x01\x02\xff\xfe\n")

    print(f"Created: {pdf_path}")


if __name__ == "__main__":
    print("Generating test CV fixtures...")

    try:
        generate_sample_cv_pdf()
        generate_minimal_cv_pdf()
        generate_sample_cv_docx()
        generate_malformed_pdf()

        print("\nAll fixtures generated successfully!")
        print("Run tests with: pytest backend/tests/test_document_processor.py")

    except Exception as e:
        print(f"Error generating fixtures: {e}")
        print("Install dependencies: pip install reportlab python-docx")
